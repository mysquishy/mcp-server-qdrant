"""Content processor for handling different file types in Qdrant MCP Server."""
import os
import io
from typing import Any, Dict, List, Optional, Tuple, Union
import logging
import json
from datetime import datetime
import re
from urllib.parse import urlparse
import mimetypes
import tempfile

from mcp.server.fastmcp import Context

logger = logging.getLogger(__name__)

try:
    import httpx
    from bs4 import BeautifulSoup
    import chardet
    WEB_DEPENDENCIES_AVAILABLE = True
except ImportError:
    logger.warning("Web dependencies not installed. Install with: pip install httpx beautifulsoup4 chardet")
    WEB_DEPENDENCIES_AVAILABLE = False

# Try to import document processing libraries
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    logger.warning("PDF processing not available. Install with: pip install PyPDF2")
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    logger.warning("DOCX processing not available. Install with: pip install python-docx")
    DOCX_AVAILABLE = False

try:
    import openpyxl
    XLSX_AVAILABLE = True
except ImportError:
    logger.warning("XLSX processing not available. Install with: pip install openpyxl")
    XLSX_AVAILABLE = False

try:
    import csv
    CSV_AVAILABLE = True
except ImportError:
    logger.warning("CSV processing not available")
    CSV_AVAILABLE = False

# Import chunk and process for storing in Qdrant
from ..data_processing.chunk_and_process import chunk_and_process


async def process_content(
    source: str,
    content_type: Optional[str] = None,
    collection: Optional[str] = None,
    extract_metadata: bool = True,
    store_in_qdrant: bool = True,
    raw_content: Optional[bytes] = None,
    remove_html_tags: bool = True,
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    timeout: int = 30,
    extract_tables: bool = False,
    extract_images: bool = False,
    ctx: Optional[Context] = None
) -> Dict[str, Any]:
    """
    Process content from various sources and file types.

    Args:
        source: URL or file path of the content
        content_type: MIME type of the content (auto-detected if None)
        collection: Collection to store the extracted content
        extract_metadata: Whether to extract metadata
        store_in_qdrant: Whether to store in Qdrant
        raw_content: Raw content bytes (if already fetched)
        remove_html_tags: Whether to remove HTML tags from HTML content
        chunk_size: Size of each chunk when storing in Qdrant
        chunk_overlap: Overlap between chunks
        timeout: Timeout in seconds for HTTP requests
        extract_tables: Whether to extract tables from documents
        extract_images: Whether to extract images from documents
        ctx: Optional MCP context

    Returns:
        Dictionary containing the processed content and metadata
    """
    if not WEB_DEPENDENCIES_AVAILABLE:
        error_msg = "Web dependencies not installed. Install with: pip install httpx beautifulsoup4 chardet"
        if ctx:
            ctx.error(error_msg)
        return {"error": error_msg}
    
    # Check if source is a URL or file path
    is_url = source.startswith(('http://', 'https://', 'ftp://'))
    
    if ctx:
        ctx.info(f"Processing content from: {source}")
        if content_type:
            ctx.info(f"Content type specified as: {content_type}")
    
    # Fetch content if needed
    if raw_content is None:
        if is_url:
            try:
                async with httpx.AsyncClient(timeout=timeout) as client:
                    if ctx:
                        ctx.info(f"Fetching content from URL: {source}")
                    
                    response = await client.get(source, follow_redirects=True)
                    response.raise_for_status()
                    
                    raw_content = response.content
                    
                    # Get content type from response if not specified
                    if not content_type and 'content-type' in response.headers:
                        content_type = response.headers['content-type'].split(';')[0].strip()
                        if ctx:
                            ctx.info(f"Detected content type from HTTP headers: {content_type}")
            
            except httpx.RequestError as e:
                error_msg = f"Request error: {str(e)}"
                logger.error(error_msg)
                if ctx:
                    ctx.error(error_msg)
                return {"error": error_msg}
            
            except httpx.HTTPStatusError as e:
                error_msg = f"HTTP error: {e.response.status_code} {e.response.reason_phrase}"
                logger.error(error_msg)
                if ctx:
                    ctx.error(error_msg)
                return {"error": error_msg}
        
        else:
            # Local file
            try:
                if ctx:
                    ctx.info(f"Reading content from file: {source}")
                
                with open(source, 'rb') as f:
                    raw_content = f.read()
            
            except Exception as e:
                error_msg = f"Error reading file {source}: {str(e)}"
                logger.error(error_msg)
                if ctx:
                    ctx.error(error_msg)
                return {"error": error_msg}
    
    # Determine content type if not specified
    if not content_type:
        # Try to guess from file extension
        content_type, _ = mimetypes.guess_type(source)
        
        if not content_type:
            # Try to guess from content
            if raw_content.startswith(b'%PDF-'):
                content_type = 'application/pdf'
            elif raw_content.startswith(b'PK\x03\x04'):
                # Office documents are zip files, check for specific types
                if any(x in source.lower() for x in ['.docx', '.doc']):
                    content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                elif any(x in source.lower() for x in ['.xlsx', '.xls']):
                    content_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                elif any(x in source.lower() for x in ['.pptx', '.ppt']):
                    content_type = 'application/vnd.openxmlformats-officedocument.presentationml.presentation'
                else:
                    content_type = 'application/zip'
            elif raw_content.startswith(b'<!DOCTYPE html>') or raw_content.startswith(b'<html'):
                content_type = 'text/html'
            elif raw_content.strip().startswith(b'{') and raw_content.strip().endswith(b'}'):
                content_type = 'application/json'
            elif all(c < 128 for c in raw_content[:1000]):
                # Probably a text file
                content_type = 'text/plain'
            else:
                content_type = 'application/octet-stream'
        
        if ctx:
            ctx.info(f"Detected content type: {content_type}")
    
    # Initialize metadata
    metadata = {}
    if extract_metadata:
        # Basic metadata
        metadata['source'] = source
        metadata['content_type'] = content_type
        metadata['size_bytes'] = len(raw_content)
        metadata['processed_at'] = datetime.now().isoformat()
        
        if is_url:
            parsed_url = urlparse(source)
            metadata['domain'] = parsed_url.netloc
            metadata['path'] = parsed_url.path
    
    # Process content based on type
    extracted_content = ""
    tables = []
    images = []
    
    # HTML
    if content_type == 'text/html':
        if ctx:
            ctx.info("Processing HTML content")
        
        try:
            # Detect encoding
            encoding_result = chardet.detect(raw_content)
            encoding = encoding_result['encoding'] or 'utf-8'
            
            # Parse HTML
            soup = BeautifulSoup(raw_content.decode(encoding, errors='replace'), 'lxml')
            
            # Extract metadata from HTML
            if extract_metadata:
                title_tag = soup.find('title')
                metadata['title'] = title_tag.text.strip() if title_tag else "No title found"
                
                # Meta description
                description_tag = soup.find('meta', attrs={'name': 'description'})
                if description_tag and 'content' in description_tag.attrs:
                    metadata['description'] = description_tag['content']
                
                # Meta keywords
                keywords_tag = soup.find('meta', attrs={'name': 'keywords'})
                if keywords_tag and 'content' in keywords_tag.attrs:
                    metadata['keywords'] = keywords_tag['content']
                
                # Author
                author_tag = soup.find('meta', attrs={'name': 'author'})
                if author_tag and 'content' in author_tag.attrs:
                    metadata['author'] = author_tag['content']
            
            # Extract main content
            main_content = None
            for selector in [
                'article', 'main', '.content', '#content', '.post', '.article',
                '.entry-content', '.post-content', '.article-content'
            ]:
                main_content = soup.select_one(selector)
                if main_content:
                    break
            
            if main_content:
                if remove_html_tags:
                    extracted_content = main_content.get_text(separator='\n', strip=True)
                else:
                    extracted_content = str(main_content)
            else:
                # No specific content container found, extract from body
                body = soup.find('body')
                if body:
                    if remove_html_tags:
                        # Remove script, style, and other non-content elements
                        for element in body(['script', 'style', 'head', 'header', 'footer', 'nav']):
                            element.extract()
                        extracted_content = body.get_text(separator='\n', strip=True)
                    else:
                        extracted_content = str(body)
            
            # Extract tables if requested
            if extract_tables:
                table_tags = soup.find_all('table')
                for i, table_tag in enumerate(table_tags):
                    if ctx and i == 0:
                        ctx.info(f"Extracting {len(table_tags)} tables from HTML")
                    
                    # Convert table to text representation
                    table_text = ""
                    rows = table_tag.find_all('tr')
                    
                    for row in rows:
                        cells = row.find_all(['td', 'th'])
                        row_text = ' | '.join(cell.get_text(strip=True) for cell in cells)
                        table_text += row_text + '\n'
                    
                    tables.append({
                        'table_index': i,
                        'table_text': table_text
                    })
            
            # Clean up content
            if remove_html_tags:
                # Remove multiple newlines
                extracted_content = re.sub(r'\n{3,}', '\n\n', extracted_content)
                # Remove multiple spaces
                extracted_content = re.sub(r' {2,}', ' ', extracted_content)
        
        except Exception as e:
            error_msg = f"Error processing HTML content: {str(e)}"
            logger.error(error_msg)
            if ctx:
                ctx.error(error_msg)
            return {"error": error_msg}
    
    # PDF
    elif content_type == 'application/pdf':
        if not PDF_AVAILABLE:
            error_msg = "PDF processing not available. Install with: pip install PyPDF2"
            if ctx:
                ctx.error(error_msg)
            return {"error": error_msg}
        
        if ctx:
            ctx.info("Processing PDF content")
        
        try:
            # Save raw content to a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                temp_file.write(raw_content)
                temp_file_path = temp_file.name
            
            try:
                # Extract text and metadata from PDF
                with open(temp_file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    
                    # Extract metadata
                    if extract_metadata:
                        if reader.metadata:
                            for key, value in reader.metadata.items():
                                if key.startswith('/'):
                                    clean_key = key[1:].lower()
                                    metadata[clean_key] = value
                        
                        metadata['page_count'] = len(reader.pages)
                    
                    # Extract text from all pages
                    all_text = []
                    for i, page in enumerate(reader.pages):
                        if ctx and i % 10 == 0:
                            ctx.report_progress(i, len(reader.pages))
                        
                        text = page.extract_text()
                        if text:
                            all_text.append(text)
                    
                    extracted_content = '\n\n'.join(all_text)
                    
                    # Extract tables if requested (simplified - just identify table-like content)
                    if extract_tables:
                        # This is a simplified approach - proper table extraction from PDFs
                        # would require more advanced libraries like Camelot or Tabula
                        table_pattern = re.compile(r'((?:[^\n]+\|[^\n]+(?:\n|$))+)')
                        table_matches = table_pattern.findall(extracted_content)
                        
                        for i, table_text in enumerate(table_matches):
                            tables.append({
                                'table_index': i,
                                'table_text': table_text
                            })
                        
                        if ctx and tables:
                            ctx.info(f"Identified {len(tables)} potential tables in PDF")
            
            finally:
                # Remove temporary file
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
        
        except Exception as e:
            error_msg = f"Error processing PDF content: {str(e)}"
            logger.error(error_msg)
            if ctx:
                ctx.error(error_msg)
            return {"error": error_msg}
    
    # Word Document (DOCX)
    elif content_type.startswith('application/vnd.openxmlformats-officedocument.wordprocessingml.document'):
        if not DOCX_AVAILABLE:
            error_msg = "DOCX processing not available. Install with: pip install python-docx"
            if ctx:
                ctx.error(error_msg)
            return {"error": error_msg}
        
        if ctx:
            ctx.info("Processing DOCX content")
        
        try:
            # Save raw content to a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(raw_content)
                temp_file_path = temp_file.name
            
            try:
                # Extract text and metadata from DOCX
                doc = Document(temp_file_path)
                
                # Extract metadata from core properties
                if extract_metadata and hasattr(doc, 'core_properties'):
                    props = doc.core_properties
                    if hasattr(props, 'title') and props.title:
                        metadata['title'] = props.title
                    if hasattr(props, 'author') and props.author:
                        metadata['author'] = props.author
                    if hasattr(props, 'subject') and props.subject:
                        metadata['subject'] = props.subject
                    if hasattr(props, 'keywords') and props.keywords:
                        metadata['keywords'] = props.keywords
                    if hasattr(props, 'created') and props.created:
                        metadata['created'] = props.created.isoformat()
                    if hasattr(props, 'modified') and props.modified:
                        metadata['modified'] = props.modified.isoformat()
                
                # Extract text from paragraphs
                paragraphs = [p.text for p in doc.paragraphs if p.text]
                extracted_content = '\n'.join(paragraphs)
                
                # Extract tables if requested
                if extract_tables:
                    for i, table in enumerate(doc.tables):
                        table_text = ""
                        for row in table.rows:
                            row_text = ' | '.join(cell.text for cell in row.cells)
                            table_text += row_text + '\n'
                        
                        tables.append({
                            'table_index': i,
                            'table_text': table_text
                        })
                    
                    if ctx and tables:
                        ctx.info(f"Extracted {len(tables)} tables from DOCX")
            
            finally:
                # Remove temporary file
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
        
        except Exception as e:
            error_msg = f"Error processing DOCX content: {str(e)}"
            logger.error(error_msg)
            if ctx:
                ctx.error(error_msg)
            return {"error": error_msg}
    
    # Excel Spreadsheet (XLSX)
    elif content_type.startswith('application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'):
        if not XLSX_AVAILABLE:
            error_msg = "XLSX processing not available. Install with: pip install openpyxl"
            if ctx:
                ctx.error(error_msg)
            return {"error": error_msg}
        
        if ctx:
            ctx.info("Processing XLSX content")
        
        try:
            # Save raw content to a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as temp_file:
                temp_file.write(raw_content)
                temp_file_path = temp_file.name
            
            try:
                # Extract data from Excel file
                workbook = openpyxl.load_workbook(temp_file_path, data_only=True)
                
                # Extract metadata
                if extract_metadata:
                    metadata['sheet_names'] = workbook.sheetnames
                    metadata['sheet_count'] = len(workbook.sheetnames)
                
                # Extract content from all sheets
                all_sheet_data = []
                
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    sheet_data = f"Sheet: {sheet_name}\n"
                    
                    # Find the data bounds in the sheet
                    min_row, min_col, max_row, max_col = 1, 1, 1, 1
                    for row in sheet.iter_rows():
                        for cell in row:
                            if cell.value is not None:
                                max_row = max(max_row, cell.row)
                                max_col = max(max_col, cell.column)
                    
                    # Extract data within bounds
                    for row in range(1, max_row + 1):
                        row_data = []
                        for col in range(1, max_col + 1):
                            cell = sheet.cell(row=row, column=col)
                            row_data.append(str(cell.value) if cell.value is not None else "")
                        
                        sheet_data += " | ".join(row_data) + "\n"
                    
                    all_sheet_data.append(sheet_data)
                    
                    # Add to tables if extract_tables is True
                    if extract_tables:
                        tables.append({
                            'table_index': len(tables),
                            'sheet_name': sheet_name,
                            'table_text': sheet_data
                        })
                
                extracted_content = "\n\n".join(all_sheet_data)
            
            finally:
                # Remove temporary file
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
        
        except Exception as e:
            error_msg = f"Error processing XLSX content: {str(e)}"
            logger.error(error_msg)
            if ctx:
                ctx.error(error_msg)
            return {"error": error_msg}
    
    # CSV
    elif content_type == 'text/csv' or source.lower().endswith('.csv'):
        if not CSV_AVAILABLE:
            error_msg = "CSV processing not available"
            if ctx:
                ctx.error(error_msg)
            return {"error": error_msg}
        
        if ctx:
            ctx.info("Processing CSV content")
        
        try:
            # Detect encoding
            encoding_result = chardet.detect(raw_content)
            encoding = encoding_result['encoding'] or 'utf-8'
            
            # Parse CSV
            text_content = raw_content.decode(encoding, errors='replace')
            csv_reader = csv.reader(text_content.splitlines())
            rows = list(csv_reader)
            
            # Extract metadata
            if extract_metadata:
                metadata['row_count'] = len(rows)
                metadata['column_count'] = len(rows[0]) if rows else 0
                if rows and len(rows) > 0:
                    metadata['headers'] = rows[0]
            
            # Convert to text representation
            csv_text = "\n".join([" | ".join(row) for row in rows])
            extracted_content = csv_text
            
            # Add to tables if extract_tables is True
            if extract_tables:
                tables.append({
                    'table_index': 0,
                    'table_text': csv_text
                })
        
        except Exception as e:
            error_msg = f"Error processing CSV content: {str(e)}"
            logger.error(error_msg)
            if ctx:
                ctx.error(error_msg)
            return {"error": error_msg}
    
    # JSON
    elif content_type == 'application/json':
        if ctx:
            ctx.info("Processing JSON content")
        
        try:
            # Detect encoding
            encoding_result = chardet.detect(raw_content)
            encoding = encoding_result['encoding'] or 'utf-8'
            
            # Parse JSON
            text_content = raw_content.decode(encoding, errors='replace')
            json_data = json.loads(text_content)
            
            # Pretty print JSON
            extracted_content = json.dumps(json_data, indent=2)
            
            # Extract metadata
            if extract_metadata:
                if isinstance(json_data, dict):
                    metadata['root_keys'] = list(json_data.keys())
                elif isinstance(json_data, list):
                    metadata['array_length'] = len(json_data)
                    if json_data and isinstance(json_data[0], dict):
                        metadata['item_keys'] = list(json_data[0].keys())
        
        except Exception as e:
            error_msg = f"Error processing JSON content: {str(e)}"
            logger.error(error_msg)
            if ctx:
                ctx.error(error_msg)
            return {"error": error_msg}
    
    # Plain text
    elif content_type == 'text/plain':
        if ctx:
            ctx.info("Processing plain text content")
        
        try:
            # Detect encoding
            encoding_result = chardet.detect(raw_content)
            encoding = encoding_result['encoding'] or 'utf-8'
            
            # Decode text
            extracted_content = raw_content.decode(encoding, errors='replace')
            
            # Extract metadata
            if extract_metadata:
                metadata['line_count'] = extracted_content.count('\n') + 1
                metadata['word_count'] = len(extracted_content.split())
        
        except Exception as e:
            error_msg = f"Error processing text content: {str(e)}"
            logger.error(error_msg)
            if ctx:
                ctx.error(error_msg)
            return {"error": error_msg}
    
    # Unsupported format
    else:
        error_msg = f"Unsupported content type: {content_type}"
        if ctx:
            ctx.warning(error_msg)
        
        # Try to decode as text anyway
        try:
            encoding_result = chardet.detect(raw_content)
            encoding = encoding_result['encoding'] or 'utf-8'
            extracted_content = raw_content.decode(encoding, errors='replace')
        except:
            extracted_content = "[Binary content - cannot extract text]"
    
    # Store in Qdrant if requested
    points_stored = 0
    if store_in_qdrant and collection and extracted_content:
        if ctx:
            ctx.info(f"Storing content in Qdrant collection: {collection}")
        
        # Prepare metadata for storage
        storage_metadata = {
            **metadata,
            "tables_count": len(tables),
            "images_count": len(images)
        }
        
        # Add tables summary to metadata if there aren't too many
        if tables and len(tables) <= 5:
            tables_summary = []
            for table in tables:
                # Limit table text to prevent metadata size issues
                table_text = table['table_text']
                if len(table_text) > 500:
                    table_text = table_text[:497] + "..."
                
                tables_summary.append({
                    'table_index': table['table_index'],
                    'table_summary': table_text
                })
            
            storage_metadata["tables_summary"] = tables_summary
        
        # Use the chunk_and_process tool to store in Qdrant
        try:
            result = await chunk_and_process(
                text=extracted_content,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                collection=collection,
                metadata=storage_metadata,
                ctx=ctx
            )
            
            if isinstance(result, dict) and "chunks" in result:
                points_stored = len(result["chunks"])
                if ctx:
                    ctx.info(f"Successfully stored {points_stored} chunks in Qdrant")
        except Exception as e:
            error_msg = f"Error storing content in Qdrant: {str(e)}"
            logger.error(error_msg)
            if ctx:
                ctx.error(error_msg)
    
    # Return results
    return {
        "source": source,
        "content_type": content_type,
        "content_length": len(extracted_content),
        "metadata": metadata,
        "tables_count": len(tables),
        "images_count": len(images),
        "chunks_stored": points_stored,
        "collection": collection if store_in_qdrant else None,
        "content_preview": extracted_content[:500] + "..." if len(extracted_content) > 500 else extracted_content
    }
